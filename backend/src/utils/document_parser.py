# -*- coding: utf-8 -*-
"""
Document parsing utilities for PDF, Word, images, and plain text.
Optional dependencies are loaded defensively so the backend can still start
in lightweight local environments.
"""

from __future__ import annotations

import io
from typing import Optional

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import magic  # type: ignore
except ImportError:
    magic = None

try:
    from paddleocr import PaddleOCR
except ImportError:
    PaddleOCR = None


class FileTypeDetector:
    """Detect supported file types and validate file size."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "image/jpeg": "image",
        "image/png": "image",
        "text/plain": "text",
    }

    EXTENSION_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".txt": "text",
    }

    MAX_FILE_SIZE = 50 * 1024 * 1024

    @classmethod
    def detect(cls, content: bytes, filename: str) -> str:
        mime_type = cls._detect_mime_type(content, filename)
        file_type = cls.SUPPORTED_TYPES.get(mime_type)
        if file_type:
            return file_type

        fallback_type = cls._detect_from_filename(filename)
        if fallback_type:
            return fallback_type

        raise ValueError(f"不支持的文件类型: {mime_type}")

    @classmethod
    def validate_size(cls, size: int) -> None:
        if size > cls.MAX_FILE_SIZE:
            raise ValueError(
                f"文件大小超过限制 ({cls.MAX_FILE_SIZE / 1024 / 1024:.0f}MB)"
            )

    @classmethod
    def _detect_mime_type(cls, content: bytes, filename: str) -> str:
        if magic is not None:
            return magic.from_buffer(content, mime=True)

        if content.startswith(b"%PDF"):
            return "application/pdf"

        if content.startswith(b"\x89PNG"):
            return "image/png"

        if content.startswith(b"\xff\xd8\xff"):
            return "image/jpeg"

        if content.startswith(b"PK") and filename.lower().endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        if filename.lower().endswith(".txt"):
            return "text/plain"

        return "application/octet-stream"

    @classmethod
    def _detect_from_filename(cls, filename: str) -> Optional[str]:
        lower_name = filename.lower()
        for extension, file_type in cls.EXTENSION_TYPES.items():
            if lower_name.endswith(extension):
                return file_type
        return None


class PDFParser:
    """Parse text and tables from PDF documents."""

    @staticmethod
    def parse(content: bytes) -> str:
        if pdfplumber is None:
            raise ValueError("PDF 解析依赖未安装，请先安装 pdfplumber。")

        text_parts = []

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    for table in page.extract_tables():
                        table_text = "\n".join(
                            " | ".join(str(cell) if cell else "" for cell in row)
                            for row in table
                        )
                        text_parts.append(table_text)

            return "\n\n".join(text_parts)
        except Exception as exc:
            raise ValueError(f"PDF 解析失败: {exc}") from exc


class DocxParser:
    """Parse paragraphs and tables from Word documents."""

    @staticmethod
    def parse(content: bytes) -> str:
        if Document is None:
            raise ValueError("Word 解析依赖未安装，请先安装 python-docx。")

        text_parts = []

        try:
            doc = Document(io.BytesIO(content))

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    text_parts.append(" | ".join(cell.text for cell in row.cells))

            return "\n\n".join(text_parts)
        except Exception as exc:
            raise ValueError(f"Word 文档解析失败: {exc}") from exc


class ImageOCR:
    """OCR wrapper for image uploads."""

    def __init__(self):
        if PaddleOCR is None:
            self.ocr = None
            return

        try:
            self.ocr = PaddleOCR(
                use_angle_cls=True,
                lang="ch",
                use_gpu=False,
                show_log=False,
            )
        except Exception:
            self.ocr = None

    def parse(self, content: bytes) -> str:
        if not self.ocr:
            raise ValueError("OCR 引擎未初始化，请先安装图片识别依赖。")

        if Image is None:
            raise ValueError("图片解析依赖未安装，请先安装 Pillow。")

        try:
            image = Image.open(io.BytesIO(content))
            if image.mode == "RGBA":
                image = image.convert("RGB")

            result = self.ocr.ocr(image, cls=True)

            text_parts = []
            if result and result[0]:
                for line in result[0]:
                    if line and len(line) >= 2:
                        text_parts.append(line[1][0])

            return "\n".join(text_parts)
        except Exception as exc:
            raise ValueError(f"图片 OCR 识别失败: {exc}") from exc


class DocumentParser:
    """Unified parser for all supported file types."""

    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DocxParser()
        self.image_ocr = ImageOCR()

    def parse(self, content: bytes, file_type: str) -> str:
        if file_type == "pdf":
            return self.pdf_parser.parse(content)

        if file_type == "docx":
            return self.docx_parser.parse(content)

        if file_type == "image":
            return self.image_ocr.parse(content)

        if file_type == "text":
            for encoding in ("utf-8", "gbk", "gb2312"):
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法解码文本文件")

        raise ValueError(f"不支持的文件类型: {file_type}")


_parser_instance: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()
    return _parser_instance
